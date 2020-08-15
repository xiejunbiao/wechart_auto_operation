from docx import Document
from util.get_data_from_sql import GetData
doc = Document()
doc.add_heading('title1', level=1)
doc.add_heading('title1', level=2)
doc.add_heading('title1', level=3)
doc.save("first_word.docx")


class CreateDocEmpty(GetData):
    """

    """
    def __init__(self):
        GetData.__init__(self)
